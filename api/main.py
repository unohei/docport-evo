# 変更点（v1.9 OCR を画像OCR専用に刷新）:
# 1. pypdf を撤廃し pypdfium2 でPDFをページ画像化（最大3ページ、10MB上限）
# 2. OCR は OpenAI Vision API（gpt-4o）を使用（未設定なら500）
#    ※ Gemini は後続バージョンで削除済み
# 3. /api/ocr の認証: JWT + profiles テーブルで hospital_id 確認
#    ※ 送信前ファイル専用のため documents テーブル照合は行わない（v1.9 当初の実装意図と異なる）
#
# 変更点（v2.0 OCRテキストの構造化JSON生成を追加）:
# 1. _structure_referral_text: OCR済みテキストを gpt-4o で医療紹介状の構造化JSONへ変換
# 2. /api/ocr のレスポンスに structured フィールドを追加（失敗時は null、OCR全体は落とさない）
# 3. 既存のVision OCR処理・認証ロジックは一切変更なし
#
# 変更点（v2.1 presign-download の拡張子バリデーション修正）:
# 1. _assert_download_access の file_key 検証を ".pdf" ハードコードから ALLOWED_MIME_EXT 値セットに変更
#    → xlsx / docx / pptx / png / jpg など PDF 以外のファイルも presign-download できるように修正
#
# 変更点（v2.2 DOCX テキスト抽出を /api/ocr に統合）:
# 1. _extract_docx_text: python-docx でバイト列から本文テキストを抽出（Vision OCR 不使用）
# 2. _ocr_impl: file_key の拡張子が .docx の場合は DOCX 抽出ルートへ分岐
# 3. meta に source_type: "pdf"|"docx" を追加
# 4. structured / alerts は PDF OCR と同じ後段処理を流用
# 5. DOCX 抽出失敗時は graceful degradation（text="" + warnings に理由を追加、500 にしない）
#
# 変更点（v2.3 XLSX テキスト抽出を /api/ocr に統合）:
# 1. _extract_xlsx_text: openpyxl で全シートを走査しテキスト化（"Sheet:<名>" + "A:値|B:値..." 形式）
# 2. 連続空行 _XLSX_MAX_EMPTY_ROWS 超でシート打ち切り、全体 _XLSX_MAX_CHARS 超で切り詰め+警告
# 3. _ocr_impl: .xlsx 拡張子を許可し XLSX 抽出ルートへ分岐
# 4. source_type: "xlsx" を meta に追加、structured/alerts は既存後段処理を流用
#
# 変更点（v2.4 text_normalized（AI投入用正規化テキスト）を追加）:
# 1. _normalize_text: raw から AI 投入用テキストを生成（表示用 raw は変更しない）
#    (A) コードフェンス除去  (B) セル形式整形  (C) 見出し+次行結合
#    (D) 残存セル接頭辞除去  (E) 連続空行整理  (F) 8000文字上限
# 2. _ocr_impl: structured / alerts / 要配慮キーワードを text_normalized を入力に変更
# 3. レスポンスに text_normalized を追加（text=raw は維持）
#
# 変更点（v2.4.1 _normalize_text ルールC の安全ガード追加）:
# 1. 次行が _HEADING_KEYWORDS に含まれる場合は結合しない（別見出しへの誤結合を防止）
# 2. 変更は _normalize_text の条件1行のみ（他ルール・認証・RLS は無変更）
#
# 変更点（v2.4.2 _normalize_text ルールC の見出し判定を強化）:
# 1. 見出し判定: 全角/半角スペース除去後に _HEADING_KEYWORDS 完全一致、行長<15、":" "：" なし
# 2. 次行判定: 非空・長さ>1・":" "：" なし・見出しでない の4条件に強化
# 3. 既に "主訴:" 形式の行は ":" チェックにより自動スキップ（二重処理なし）
#
# 変更点（v2.5 見出し正規化キーを NFKC+ゼロ幅文字除去で強化 + debug モード追加）:
# 1. _norm_heading_key: NFKC正規化→strip→空白/ゼロ幅文字除去で文字コード揺れを吸収
# 2. _HEADINGS_NORM: _norm_heading_key で正規化済みのキー集合（判定に使用）
# 3. _normalize_text: debug=True のとき heading_matches/joined_pairs を返す（本番影響なし）
# 4. OcrRequest mode="debug": full と同等処理 + debug_normalize をレスポンスに追加
#
# 変更点（v2.5.1 ルールCのバグ修正: "A:主訴" 形式でも見出し結合が発動するよう修正）:
# 根本原因: (B)セル整形→(C)見出し結合→(D)接頭辞除去 の順で実行するため、
#           (C) 時点では "A:主訴" のまま → ":" を含む → 見出し判定が常にスキップされていた。
# 修正: ルールC内で _CELL_PREFIX_RE を先に適用した heading_text / next_body で判定・出力する。
#       処理順（B→C→D）・他ルール・認証・RLS は変更なし。

import base64
import io
import json
import logging
import os
import re
import time
import unicodedata
import urllib.error
import urllib.parse
import urllib.request
import uuid

import pypdfium2 as pdfium
from jose import jwt as jose_jwt
from jose.exceptions import ExpiredSignatureError, JWTClaimsError, JWTError

from fastapi import BackgroundTasks, Body, Depends, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel

from r2_client import get_bucket_name, get_s3_client

from typing import Optional, List, Dict, Tuple

app = FastAPI()

# ----------------------------
# CORS（本番 + ローカル）
# allow_origins は使用しない。allow_origin_regex のみで制御する。
# デフォルト: docport.pages.dev / docport-evo.pages.dev / localhost:5173
# ----------------------------
_DEFAULT_ORIGIN_REGEX = (
    r"^("
    r"https://docport\.pages\.dev"
    r"|https://docport-evo\.pages\.dev"
    r"|http://localhost:5173"
    r")$"
)
ALLOW_ORIGIN_REGEX = os.getenv("ALLOW_ORIGIN_REGEX", _DEFAULT_ORIGIN_REGEX)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[],           # allow_origin_regex に一本化するため空にする
    allow_origin_regex=ALLOW_ORIGIN_REGEX,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ----------------------------
# Supabase 設定
# ----------------------------
SUPABASE_URL = os.getenv("SUPABASE_URL", "")        # 例: https://xxxx.supabase.co
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", "")

_bearer = HTTPBearer()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)


# ----------------------------
# JWKS キャッシュ（TTL: 1時間）
# ----------------------------
_jwks_keys: dict = {}     # kid -> JWK dict
_jwks_fetched_at: float = 0.0
_JWKS_CACHE_TTL = 3600    # seconds


def _refresh_jwks() -> None:
    """
    Supabase の JWKS エンドポイントから公開鍵を取得してモジュール変数に格納する。
    urllib のみ使用（新規ライブラリ不要）。
    """
    global _jwks_keys, _jwks_fetched_at
    if not SUPABASE_URL:
        raise HTTPException(status_code=500, detail="SUPABASE_URL が未設定です")
    url = f"{SUPABASE_URL.rstrip('/')}/auth/v1/.well-known/jwks.json"
    req = urllib.request.Request(url, headers={"Accept": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        logger.error("JWKS 取得失敗 (%d)", e.code)
        raise HTTPException(status_code=500, detail="認証サービスへの接続に失敗しました")
    except Exception:
        logger.exception("JWKS 取得エラー")
        raise HTTPException(status_code=500, detail="認証サービスへの接続に失敗しました")

    _jwks_keys = {k["kid"]: k for k in data.get("keys", []) if "kid" in k}
    _jwks_fetched_at = time.time()


def _get_signing_key(kid: str) -> dict:
    """
    kid に対応する JWK dict を返す（TTL キャッシュ付き）。
    キャッシュにない場合は強制再取得（key rotation 対応）。
    """
    global _jwks_keys, _jwks_fetched_at

    # TTL 切れなら再取得
    if time.time() - _jwks_fetched_at > _JWKS_CACHE_TTL:
        _refresh_jwks()

    if kid in _jwks_keys:
        return _jwks_keys[kid]

    # キャッシュにない → key rotation の可能性があるので強制再取得
    _refresh_jwks()
    if kid not in _jwks_keys:
        raise HTTPException(
            status_code=401,
            detail="対応する公開鍵が見つかりません（kid 不一致）",
        )
    return _jwks_keys[kid]


# ----------------------------
# JWT 検証ヘルパー（ES256/JWKS）
# ----------------------------
def verify_jwt(credentials: HTTPAuthorizationCredentials = Depends(_bearer)) -> dict:
    """
    Supabase JWT（ES256）を JWKS 経由で検証する。
    - kid から公開鍵を取得（TTL キャッシュ付き）
    - audience: authenticated
    - issuer: https://<SUPABASE_URL>/auth/v1
    FastAPI の依存注入で同一リクエスト内は _bearer の結果がキャッシュされる。
    """
    token = credentials.credentials

    # ヘッダーから kid / alg を取得（未検証）
    try:
        header = jose_jwt.get_unverified_header(token)
    except JWTError:
        raise HTTPException(status_code=401, detail="無効なトークンです（ヘッダー解析失敗）")

    kid = header.get("kid")
    alg = header.get("alg", "ES256")

    if not kid:
        raise HTTPException(status_code=401, detail="JWT に kid がありません")

    jwk_key = _get_signing_key(kid)
    issuer = f"{SUPABASE_URL.rstrip('/')}/auth/v1"

    try:
        # python-jose は {"keys": [...]} 形式の JWKS dict を直接受け取れる
        payload = jose_jwt.decode(
            token,
            {"keys": [jwk_key]},
            algorithms=[alg],
            audience="authenticated",
            issuer=issuer,
        )
        return payload
    except ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="トークンの有効期限が切れています")
    except JWTClaimsError:
        raise HTTPException(status_code=401, detail="トークンのクレームが無効です")
    except JWTError:
        raise HTTPException(status_code=401, detail="無効なトークンです")


# ----------------------------
# Supabase REST API ヘルパー
# ----------------------------
def _supabase_get(path: str, jwt_token: str) -> list:
    """
    user JWT を使って Supabase REST API を GET する（RLS が有効に機能する）。
    service_role キーは使わない。
    """
    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_ANON_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    req = urllib.request.Request(
        url,
        headers={
            "apikey": SUPABASE_ANON_KEY,
            "Authorization": f"Bearer {jwt_token}",
            "Accept": "application/json",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=5) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        logger.error("Supabase user GET エラー (%d): path=%s", e.code, path)
        raise HTTPException(status_code=502, detail="データベース接続エラーが発生しました")
    except Exception:
        logger.exception("Supabase user GET 接続エラー: path=%s", path)
        raise HTTPException(status_code=502, detail="データベース接続エラーが発生しました")


def _get_hospital_id(user_id: str, jwt_token: str) -> str:
    """profiles テーブルから呼び出し元ユーザーの hospital_id を取得する"""
    uid_encoded = urllib.parse.quote(user_id, safe="")
    rows = _supabase_get(
        f"profiles?id=eq.{uid_encoded}&select=hospital_id",
        jwt_token,
    )
    if not rows or not rows[0].get("hospital_id"):
        raise HTTPException(
            status_code=403,
            detail="プロフィールが見つかりません（hospital_id 未設定）",
        )
    return rows[0]["hospital_id"]


def _assert_download_access(file_key: str, hospital_id: str, jwt_token: str) -> None:
    """
    documents テーブルを user JWT で照会し、
    from_hospital_id OR to_hospital_id がユーザーの病院と一致するか確認する。
    RLS でも弾かれるが、FastAPI 側でも明示チェック（二重防御）。
    """
    # パストラバーサル防止（基本バリデーション）
    # 許可拡張子は ALLOWED_MIME_EXT の値セットと一致させる
    ext = file_key.rsplit(".", 1)[-1].lower() if "." in file_key else ""
    if not file_key.startswith("documents/") or ext not in set(ALLOWED_MIME_EXT.values()):
        raise HTTPException(status_code=400, detail="無効な file_key です")

    key_encoded = urllib.parse.quote(file_key, safe="")
    rows = _supabase_get(
        f"documents?file_key=eq.{key_encoded}&select=from_hospital_id,to_hospital_id",
        jwt_token,
    )

    # RLS で弾かれた場合も rows が空になるため、存在有無を区別しない（情報漏洩防止）
    doc = rows[0] if rows else {}
    if not rows:
        raise HTTPException(status_code=403, detail="ドキュメントへのアクセス権がありません")

    if (
        doc.get("from_hospital_id") != hospital_id
        and doc.get("to_hospital_id") != hospital_id
    ):
        raise HTTPException(status_code=403, detail="ドキュメントへのアクセス権がありません")


def _assert_fax_file_key(file_key: str) -> None:
    """
    FAX送信専用のファイルアクセスチェック。
    _assert_download_access（documents レコード前提）の代替として使用する。

    チェック内容:
    1. file_key のフォーマット検証（パストラバーサル防止）
    2. R2 に実際にファイルが存在するか確認（head_object）

    設計メモ:
    - FAX送信時点では documents レコードがまだ存在しないため documents 照合はしない。
    - hospital_id との紐付けは検証できないが、file_key は presign-upload 時に
      JWT認証済みユーザーのみが取得できるランダム UUID のため MVP では許容する。
    # TODO: 将来的に temp_uploads テーブル（file_key, hospital_id, user_id, expires_at）を追加し、
    #        presign-upload 時に INSERT → ここで hospital_id 一致チェックに昇格させること。
    """
    # FAX送信は PDF のみ（CloudFAX API が PDF を前提とするため他形式は拒否）
    ext = file_key.rsplit(".", 1)[-1].lower() if "." in file_key else ""
    if not file_key.startswith("documents/") or ext != "pdf":
        raise HTTPException(status_code=400, detail="FAX送信は PDF ファイルのみ対応しています")

    try:
        bucket = get_bucket_name()
        s3     = get_s3_client()
        s3.head_object(Bucket=bucket, Key=file_key)
    except Exception as e:
        # NoSuchKey / 404 系は 403 で返す（存在確認は情報漏洩になるため区別しない）
        logger.warning("[_assert_fax_file_key] R2 head_object 失敗: file_key=%s err=%s", file_key, e)
        raise HTTPException(status_code=403, detail="ファイルが見つからないか、アクセスできません")


def _get_fax_contact(contact_id: str, hospital_id: str, jwt_token: str) -> dict:
    """
    contacts テーブルから FAX送信先を取得し、送信可否を検証する。

    チェック内容:
    1. レコード存在確認
    2. hospital_id が呼び出し元と一致すること（他院の contact は使用不可）
    3. is_active == true であること
    4. fax_number が設定されていること

    戻り値: {"id", "fax_number", "hospital_id", "is_active"} を含む dict

    # TODO(temp_uploads): 将来的に _assert_fax_file_key もここに統合し、
    #   temp_uploads.hospital_id と contacts.hospital_id の一致まで検証できるようにする。
    """
    cid_enc = urllib.parse.quote(contact_id, safe="")
    rows = _supabase_get(
        f"contacts?id=eq.{cid_enc}&select=id,fax_number,hospital_id,is_active",
        jwt_token,
    )
    if not rows:
        raise HTTPException(status_code=404, detail="FAX送信先が見つかりません")

    contact = rows[0]
    if contact.get("hospital_id") != hospital_id:
        raise HTTPException(status_code=403, detail="自院のFAX送信先のみ使用できます")

    if not contact.get("is_active"):
        raise HTTPException(status_code=400, detail="このFAX送信先は無効です（is_active=false）")

    fax_number = (contact.get("fax_number") or "").strip()
    if not fax_number:
        raise HTTPException(status_code=400, detail="FAX番号が登録されていません")

    return contact


# ----------------------------
# ヘルスチェック
# ----------------------------
@app.get("/")
def health_root():
    return {"status": "ok"}


@app.get("/health")
def health():
    """ヘルスチェック（認証不要）。cold start 対策として使用する。"""
    return {"status": "ok"}


# ----------------------------
# アップロード許可 MIME マップ（許可リスト方式・単一の真実）
# フロント → FastAPI に content_type を渡し、ここで検証してから presign を発行する。
# 拡張子だけに依存せず MIME で判定する（二重防御の最終判断はここ）
# ----------------------------
ALLOWED_MIME_EXT: dict[str, str] = {
    "application/pdf":                                                              "pdf",
    "image/png":                                                                    "png",
    "image/jpeg":                                                                   "jpg",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document":      "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":            "xlsx",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation":    "pptx",
}


class PresignUploadRequest(BaseModel):
    content_type: str = "application/pdf"   # MIME タイプ（未送信時は PDF）
    filename: str = ""                      # オリジナルファイル名（ログ用）


# ----------------------------
# Presign 内部ヘルパー
# ----------------------------
def _presign_upload(content_type: str = "application/pdf") -> dict:
    # MIME 許可リスト検証（拡張子は MIME から決定する）
    if content_type not in ALLOWED_MIME_EXT:
        raise HTTPException(
            status_code=400,
            detail=f"許可されていないファイル形式です: {content_type}。"
                   f"対応形式: {', '.join(ALLOWED_MIME_EXT.keys())}",
        )
    ext = ALLOWED_MIME_EXT[content_type]

    try:
        bucket = get_bucket_name()
        s3 = get_s3_client()
    except Exception:
        logger.exception("R2クライアント初期化失敗 (upload)")
        raise HTTPException(status_code=500, detail="ストレージ接続エラーが発生しました")

    key = f"documents/{uuid.uuid4()}.{ext}"
    url = s3.generate_presigned_url(
        ClientMethod="put_object",
        Params={
            "Bucket": bucket,
            "Key": key,
            "ContentType": content_type,   # presigned URL に ContentType を含めることで PUT 時の MIME を強制
        },
        ExpiresIn=60 * 10,  # 10分（アップロード操作）
    )
    return {"upload_url": url, "file_key": key, "content_type": content_type, "file_ext": ext}


def _presign_download(key: str):
    try:
        bucket = get_bucket_name()
        s3 = get_s3_client()
    except Exception:
        logger.exception("R2クライアント初期化失敗 (download)")
        raise HTTPException(status_code=500, detail="ストレージ接続エラーが発生しました")

    url = s3.generate_presigned_url(
        ClientMethod="get_object",
        Params={"Bucket": bucket, "Key": key},
        ExpiresIn=60 * 5,
    )
    return {"download_url": url}


# ----------------------------
# Presign エンドポイント（JWT 認可）
# ----------------------------
@app.post("/api/presign-upload")
def presign_upload_api(
    body: Optional[PresignUploadRequest] = Body(default=None),
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """ログイン済みユーザーのみ署名 URL を発行（JWT 必須）。
    body 未送信（後方互換クライアント）の場合は application/pdf として扱う。"""
    req = body if body is not None else PresignUploadRequest()
    return _presign_upload(req.content_type)


@app.get("/api/presign-download")
def presign_download_api(
    key: str,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """JWT 検証 + documents の hospital_id チェック後に署名 URL を発行"""
    jwt_token = credentials.credentials
    user_id = user.get("sub", "")
    hospital_id = _get_hospital_id(user_id, jwt_token)
    _assert_download_access(key, hospital_id, jwt_token)
    return _presign_download(key)


@app.post("/presign-upload")
def presign_upload_compat(
    body: Optional[PresignUploadRequest] = Body(default=None),
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """compat: Vite proxy 経由のローカル開発用（同じ認可）"""
    req = body if body is not None else PresignUploadRequest()
    return _presign_upload(req.content_type)


@app.get("/presign-download")
def presign_download_compat(
    key: str,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """compat: Vite proxy 経由のローカル開発用（同じ認可）"""
    jwt_token = credentials.credentials
    user_id = user.get("sub", "")
    hospital_id = _get_hospital_id(user_id, jwt_token)
    _assert_download_access(key, hospital_id, jwt_token)
    return _presign_download(key)


# ----------------------------
# OCR 設定
# ----------------------------
_MAX_PDF_SIZE_BYTES = 10 * 1024 * 1024   # 10MB（PDF / DOCX / XLSX 共通上限）
_MAX_OCR_PAGES = 3                        # PDF ページ上限
_OCR_TIMEOUT_SECS = 30                    # 処理全体のタイムアウト（秒）

_XLSX_MAX_CHARS = 20_000      # XLSX 全体テキスト上限文字数
_XLSX_MAX_EMPTY_ROWS = 30     # 連続空行がこれ以上続いたらシート打ち切り

_NORMALIZED_MAX_CHARS = 8_000   # text_normalized の最大文字数（AI投入用）

# 見出し+次行結合の対象キーワード（ルールC）
_HEADING_KEYWORDS = frozenset([
    "主訴", "現病歴", "既往歴", "内服薬", "アレルギー", "検査所見", "紹介目的",
])

# セル接頭辞パターン（行頭の "A:", "AB:", "ABC:" など）
_CELL_PREFIX_RE = re.compile(r"^[A-Z]{1,3}:")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

# 要配慮情報の注意喚起キーワード（断定ではなく可能性の通知のみ）
_SENSITIVE_KEYWORDS = [
    "病名", "診断", "障害", "検査結果", "投薬", "処方",
    "手術", "入院", "HIV", "感染症", "精神", "がん",
]

# アラートキーワード定義（severity: high / medium / low）
# evidence に前後30文字スニペット付きで返す（断定禁止トーン）
_ALERT_KEYWORDS = [
    {"id": "hiv",        "label": "HIV/AIDS",   "severity": "high",   "keywords": ["HIV", "AIDS"]},
    {"id": "mental",     "label": "精神疾患",     "severity": "high",   "keywords": ["精神", "うつ病", "統合失調"]},
    {"id": "cancer",     "label": "悪性腫瘍",     "severity": "high",   "keywords": ["がん", "癌", "腫瘍", "悪性"]},
    {"id": "diagnosis",  "label": "病名・診断",   "severity": "medium", "keywords": ["病名", "診断"]},
    {"id": "test",       "label": "検査結果",     "severity": "medium", "keywords": ["検査結果", "検査値", "検査所見"]},
    {"id": "meds",       "label": "投薬・処方",   "severity": "medium", "keywords": ["投薬", "処方", "処方薬"]},
    {"id": "infect",     "label": "感染症",       "severity": "medium", "keywords": ["感染症", "感染"]},
    {"id": "disability", "label": "障害",         "severity": "medium", "keywords": ["障害"]},
    {"id": "surgery",    "label": "手術",         "severity": "low",    "keywords": ["手術", "術後"]},
    {"id": "admit",      "label": "入院",         "severity": "low",    "keywords": ["入院"]},
]

_OCR_PROMPT = (
    "以下の医療文書の画像に含まれるテキストをすべて正確に抽出してください。"
    "レイアウトをできる限り維持し、文字を漏れなく出力してください。"
)

# ----------------------------
# 構造化 設定
# ----------------------------
_STRUCTURE_TIMEOUT_SECS = 20   # Vision OCR とは別タイムアウト（失敗しても OCR 全体は落とさない）

_STRUCTURE_PROMPT = """\
以下は医療紹介状から抽出したテキストです。
次のJSONフォーマットで情報を抽出してください。
- 不明・記載なしの項目は null にしてください
- テキストに明示されていない情報は推測しないでください
- 余計な説明文や前置きは不要です。JSONのみ出力してください

{
  "patient_name": null,
  "patient_id": null,
  "birth_date": null,
  "referrer_hospital": null,
  "referrer_doctor": null,
  "referral_to_hospital": null,
  "referral_date": null,
  "chief_complaint": null,
  "suspected_diagnosis": null,
  "allergies": null,
  "medications": null
}
"""


# ----------------------------
# OCR 内部ヘルパー
# ----------------------------
def _render_pdf_to_png_list(pdf_bytes: bytes) -> tuple[list[bytes], int]:
    """
    pypdfium2 でPDFをページ画像化する。
    - 最大 _MAX_OCR_PAGES ページまで処理
    - scale=2.0（約144 DPI）でレンダリング（OCR精度向上）
    - 戻り値: (PNG bytes のリスト, 総ページ数)
    """
    pdf = pdfium.PdfDocument(pdf_bytes)
    total_pages = len(pdf)
    pages_to_render = min(total_pages, _MAX_OCR_PAGES)

    png_list: list[bytes] = []
    for i in range(pages_to_render):
        page = pdf[i]
        bitmap = page.render(scale=2.0)
        pil_image = bitmap.to_pil()
        buf = io.BytesIO()
        pil_image.save(buf, format="PNG")
        png_list.append(buf.getvalue())

    return png_list, total_pages


def _extract_docx_text(docx_bytes: bytes) -> tuple[str, list[str]]:
    """
    python-docx で DOCX 本文テキストを抽出する。
    - 段落テキストを順に連結（表・ヘッダー・フッターは含まない）
    - 失敗時は空文字 + 警告を返す（graceful degradation: 500 にしない）
    """
    try:
        from docx import Document as DocxDocument  # noqa: PLC0415
        doc = DocxDocument(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs), []
    except Exception as e:
        return "", [f"DOCX抽出失敗: {e}。内容を確認の上、送信可否を判断してください。"]


def _extract_xlsx_text(xlsx_bytes: bytes) -> tuple[str, list[str]]:
    """
    openpyxl で XLSX の全シートをテキスト化する。
    - 各シートの先頭に "Sheet: <name>" を出力
    - 各行は空でないセルのみ "A:値 | B:値 | ..." 形式に整形
    - 連続空行が _XLSX_MAX_EMPTY_ROWS 以上続いたらそのシートを打ち切る
    - 全体テキストが _XLSX_MAX_CHARS を超えたら切り詰め、warnings に追記
    - 失敗時は空文字 + 警告を返す（graceful degradation: 500 にしない）
    """
    try:
        import openpyxl  # noqa: PLC0415
        from openpyxl.utils import get_column_letter  # noqa: PLC0415
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    except Exception as e:
        return "", [f"XLSX読み込み失敗: {e}。内容を確認の上、送信可否を判断してください。"]

    lines: list[str] = []
    truncated = False

    try:
        for sheet in wb.worksheets:
            lines.append(f"Sheet: {sheet.title}")
            empty_row_count = 0

            for row in sheet.iter_rows(values_only=True):
                # 行が全空かチェック
                if all(cell is None or str(cell).strip() == "" for cell in row):
                    empty_row_count += 1
                    if empty_row_count >= _XLSX_MAX_EMPTY_ROWS:
                        break   # このシートはここで打ち切り
                    continue
                empty_row_count = 0

                # 非空セルのみ "列名:値" にして連結
                cells = []
                for col_idx, cell in enumerate(row):
                    if cell is None or str(cell).strip() == "":
                        continue
                    cells.append(f"{get_column_letter(col_idx + 1)}:{str(cell).strip()}")
                if cells:
                    lines.append(" | ".join(cells))

            # シートを処理するたびに全体上限をチェック
            if len("\n".join(lines)) >= _XLSX_MAX_CHARS:
                truncated = True
                break

        wb.close()
    except Exception as e:
        # 途中まで抽出できたぶんは返す
        return "\n".join(lines), [
            f"XLSX抽出中にエラーが発生しました: {e}。抽出できた部分のみ表示しています。"
        ]

    text = "\n".join(lines)
    warnings: list[str] = []
    if truncated:
        text = text[:_XLSX_MAX_CHARS]
        warnings.append(
            f"XLSXのテキストが長すぎるため {_XLSX_MAX_CHARS:,} 文字で省略しました。"
            "全内容を確認の上、送信可否を判断してください。"
        )
    return text, warnings


def _call_openai_ocr(png_list: list[bytes], timeout: float) -> str:
    """
    OpenAI Vision API（gpt-4o）でページ画像をまとめてOCRする。
    全ページを1リクエストで送信してテキストを取得する。
    """
    content: list[dict] = []
    for png in png_list:
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{base64.b64encode(png).decode()}"
            },
        })
    content.append({"type": "text", "text": _OCR_PROMPT})

    payload = json.dumps({
        "model": "gpt-4o",
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 4096,
    }).encode()
    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=int(timeout)) as resp:
            data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode(errors="replace")
        logger.error("OpenAI API HTTPエラー (%d): %s", e.code, body)
        raise HTTPException(status_code=502, detail="OCR処理でエラーが発生しました")
    except Exception as e:
        logger.exception("OpenAI API 接続エラー")
        raise HTTPException(status_code=502, detail="OCR処理でエラーが発生しました")

    try:
        return data["choices"][0]["message"]["content"]
    except (KeyError, IndexError):
        logger.error("OpenAI レスポンス解析失敗: %s", data)
        raise HTTPException(status_code=502, detail="OCR処理でエラーが発生しました")


# ----------------------------
# テキスト正規化ヘルパー
# ----------------------------
def _norm_heading_key(s: str) -> str:
    """
    見出し判定用の正規化キーを生成する。
    NFKC 正規化 → strip → 全角/半角スペース・ゼロ幅文字を除去。
    文字コードの揺れ（半角カナ、結合文字、不可視文字等）を吸収する。
    """
    s = unicodedata.normalize("NFKC", s)
    s = s.strip()
    for ch in ("\u0020", "\u3000", "\u200b", "\u200c", "\u200d", "\ufeff"):
        s = s.replace(ch, "")
    return s


# NFKC 正規化済みの見出しキー集合（_norm_heading_key と同じ変換で生成）
_HEADINGS_NORM: frozenset[str] = frozenset(_norm_heading_key(h) for h in _HEADING_KEYWORDS)


def _strip_code_fences(text: str) -> str:
    """
    OCR結果がコードフェンス（```...```）で全体を囲まれている場合に本文のみを返す。
    OpenAI が Markdown 形式で応答するケースへの対処。
    例: ```plaintext\n本文\n``` → 本文
    部分的なフェンスや複数フェンスは除去しない（全体を囲む1ブロックのみ対象）。
    """
    t = text.strip()
    if not t.startswith("```"):
        return t
    first_newline = t.find("\n")
    if first_newline < 0:
        return t
    if not t.endswith("```"):
        return t
    inner = t[first_newline + 1 : len(t) - 3]
    return inner.strip()


def _normalize_text(
    raw: str,
    *,
    debug: bool = False,
) -> Tuple[str, Optional[Dict]]:
    """
    AI投入用テキストの正規化。表示用の raw は変更しない（呼び出し元で使い分ける）。

    (A) コードフェンス除去（_strip_code_fences 流用）
    (B) XLSX セル形式 "A:ラベル | B:値 | ..." → "ラベル: 値 / ..."
        条件: パイプ区切り かつ 各パートが列名プレフィクス（A:, BC: 等）で始まる
        変換: 偶数インデックス=キー, 奇数インデックス=値 のペアにまとめる
    (C) 対象見出し単独行 + 次行の結合 "主訴\n右下腹部痛" → "主訴: 右下腹部痛"
        見出し判定: _norm_heading_key で正規化後 _HEADINGS_NORM に完全一致
                    かつ 行長 < 15、":" "：" を含まない
        次行判定:   非空、長さ > 1、":" "：" を含まない、_HEADINGS_NORM に含まれない
        既に "主訴:" 形式の行は ":" チェックで自動スキップ（二重処理なし）
    (D) 残存セル接頭辞除去（行頭の "A:" "BC:" 等を除去）
    (E) 連続空行を最大2行まで
    (F) 最大 _NORMALIZED_MAX_CHARS 文字で切り詰め + "...(truncated)"

    戻り値: (normalized_text, debug_info)
            debug=True のとき debug_info に heading_matches / joined_pairs を格納。
            debug=False のとき debug_info は None（本番影響なし）。
    """
    dbg_matches: list[dict] = []
    dbg_pairs:   list[dict] = []

    # (A)
    text = _strip_code_fences(raw)

    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # (B) セル形式行の変換
        if "|" in line:
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 2 and all(_CELL_PREFIX_RE.match(p) for p in parts if p):
                vals = [_CELL_PREFIX_RE.sub("", p).strip() for p in parts if p]
                pairs: list[str] = []
                j = 0
                while j < len(vals):
                    if j + 1 < len(vals):
                        k, v = vals[j], vals[j + 1]
                        if k and v:
                            pairs.append(f"{k}: {v}")
                        elif k:
                            pairs.append(k)
                        elif v:
                            pairs.append(v)
                        j += 2
                    else:
                        if vals[j]:
                            pairs.append(vals[j])
                        j += 1
                out.append(" / ".join(pairs))
                i += 1
                continue

        # (C) 見出し単独行 + 次行の結合（強化版）
        # 【重要】XLSX では "(B)セル整形" を通過しなかった単独セル行が "A:主訴" の形で残る。
        # (D)の接頭辞除去はこの後なので、ここで _CELL_PREFIX_RE を先に剥がして判定する。
        # → "A:主訴" → heading_text="主訴" → _HEADINGS_NORM に一致 → 結合発動
        # → PDF/DOCX には "A:" 接頭辞がないため heading_text == stripped_line となり既存動作を維持
        stripped_line = line.strip()
        heading_text = _CELL_PREFIX_RE.sub("", stripped_line).strip()   # 接頭辞を先に除去
        candidate = _norm_heading_key(heading_text)
        matched = candidate in _HEADINGS_NORM

        # debug: 見出し候補行（短く ":" なし）の一致結果を記録（不可視文字の特定に使用）
        if debug and len(heading_text) < 15 and ":" not in heading_text and "：" not in heading_text:
            dbg_matches.append({
                "raw_line":    line,
                "stripped":    stripped_line,
                "heading_text": heading_text,
                "candidate":   candidate,
                "matched":     matched,
                "codepoints":  [ord(c) for c in heading_text],
            })

        if (
            matched
            and len(heading_text) < 15
            and ":" not in heading_text
            and "：" not in heading_text
            and i + 1 < len(lines)
        ):
            next_line = lines[i + 1].strip()
            next_body = _CELL_PREFIX_RE.sub("", next_line).strip()    # 次行も接頭辞を先に除去
            next_cand = _norm_heading_key(next_body)
            if (
                next_body
                and len(next_body) > 1
                and ":" not in next_body
                and "：" not in next_body
                and next_cand not in _HEADINGS_NORM
            ):
                result_line = f"{candidate}: {next_body}"
                if debug:
                    dbg_pairs.append({
                        "heading":     candidate,
                        "body":        next_body,
                        "result_line": result_line,
                    })
                out.append(result_line)
                i += 2
                continue

        out.append(line)
        i += 1

    text = "\n".join(out)

    # (D) 残存セル接頭辞除去（行頭の "A:", "BC:" 等）
    text = re.sub(r"(?m)^[A-Z]{1,3}:", "", text)

    # (E) 連続空行を最大2行まで
    text = re.sub(r"\n{3,}", "\n\n", text)

    text = text.strip()

    # (F) 最大文字数制限
    if len(text) > _NORMALIZED_MAX_CHARS:
        text = text[:_NORMALIZED_MAX_CHARS] + "...(truncated)"

    debug_info = (
        {"heading_matches": dbg_matches, "joined_pairs": dbg_pairs}
        if debug else None
    )
    return text, debug_info


# ----------------------------
# アラート生成ヘルパー
# ----------------------------
def _generate_alerts(text: str) -> list[dict]:
    """
    テキストから要配慮キーワードを検索し、注意喚起リストを返す。
    - 断定禁止：「可能性があります」「確認してください」トーンのみ
    - evidence: キーワード前後30文字のスニペット（最大3件/アラート）
    - フロントエンドでのハイライト表示を前提に keyword フィールドも返す
    """
    alerts: list[dict] = []
    for entry in _ALERT_KEYWORDS:
        evidence: list[dict] = []
        seen: set[int] = set()
        for kw in entry["keywords"]:
            pos = 0
            while len(evidence) < 3:
                idx = text.find(kw, pos)
                if idx < 0:
                    break
                if idx not in seen:
                    seen.add(idx)
                    s = max(0, idx - 30)
                    e = min(len(text), idx + len(kw) + 30)
                    snippet = text[s:e]
                    if s > 0:
                        snippet = "…" + snippet
                    if e < len(text):
                        snippet = snippet + "…"
                    evidence.append({"page": 1, "snippet": snippet, "keyword": kw})
                pos = idx + 1
        if evidence:
            alerts.append({
                "id":       entry["id"],
                "label":    entry["label"],
                "severity": entry["severity"],
                "keyword":  entry["keywords"][0],
                "evidence": evidence,
            })
    return alerts


# ----------------------------
# 構造化 内部ヘルパー
# ----------------------------
def _structure_referral_text(
    text: str,
    timeout: float = _STRUCTURE_TIMEOUT_SECS,
) -> Optional[Dict]:
    """
    OCRで抽出したテキストを OpenAI gpt-4o で医療紹介状の構造化JSONに変換する。
    - OPENAI_API_KEY 未設定、text が空、API エラーの場合はすべて None を返す
    - 失敗してもOCRレスポンス全体は落とさない（graceful degradation）
    """
    if not OPENAI_API_KEY or not text.strip():
        return None

    payload = json.dumps({
        "model": "gpt-4o",
        "temperature": 0,
        "messages": [
            {
                "role": "user",
                "content": f"{_STRUCTURE_PROMPT}\nテキスト:\n{text}",
            }
        ],
        "max_tokens": 1024,
    }).encode()

    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=int(timeout)) as resp:
            data = json.loads(resp.read())
        raw = data["choices"][0]["message"]["content"].strip()
        # Markdownコードブロック（```json...```）も含め、最初の { ～ 最後の } を抽出してパース
        start = raw.find("{")
        end = raw.rfind("}") + 1
        if start < 0 or end <= start:
            return None
        return json.loads(raw[start:end])
    except Exception:
        return None


# ----------------------------
# OCR 実装（/api/ocr と /ocr の共通処理）
# ----------------------------
class OcrRequest(BaseModel):
    file_key: str
    mode: str = "full"
    # "full"      : structured + alerts を実行（通常）
    # "text_only" : structured をスキップ
    # "debug"     : full と同等 + debug_normalize をレスポンスに追加（本番以外で使用）


def _ocr_impl(
    body: OcrRequest,
    credentials: HTTPAuthorizationCredentials,
    user: dict,
) -> dict:
    """
    PDF画像OCR / DOCX・XLSXテキスト抽出の共通実装。
    1. file_key バリデーション（.pdf / .docx / .xlsx のみ受け付ける）
    2. JWT + hospital_id 確認（送信前ファイル＝まだ documents 未登録のため DB照合はしない）
    3. R2 からファイル取得（Presigned GET）+ サイズチェック
    4a. PDF:  pypdfium2 でページ画像化 → Gemini/OpenAI Vision OCR
    4b. DOCX: python-docx でテキスト抽出（Vision API 不使用。失敗は graceful degradation）
    4c. XLSX: openpyxl で全シートをテキスト化（Vision API 不使用。失敗は graceful degradation）
    5. OpenAI gpt-4o で構造化JSON生成（失敗時は structured=null）
    6. 結果テキスト + メタ（source_type 含む）+ 警告 + 構造化JSON + アラート を返す
    """
    start_time = time.time()

    # ---- タイムアウト残時間チェック（内部ヘルパー） ----
    def _remaining() -> float:
        elapsed = time.time() - start_time
        remaining = _OCR_TIMEOUT_SECS - elapsed
        if remaining <= 2.0:
            raise HTTPException(status_code=504, detail="OCR処理がタイムアウトしました")
        return remaining

    # ---- file_key バリデーション（パストラバーサル防止） ----
    # 対応拡張子: pdf / docx
    fkey = body.file_key
    ext = fkey.rsplit(".", 1)[-1].lower() if "." in fkey else ""
    if not fkey.startswith("documents/") or ext not in {"pdf", "docx", "xlsx"}:
        raise HTTPException(status_code=400, detail="無効な file_key です（対応: .pdf / .docx / .xlsx）")

    # ---- JWT + hospital_id 確認 ----
    # OCR は「送信前」専用のため documents テーブルにまだレコードが存在しない。
    # _assert_download_access（DB照合）は行わず、有効な JWT + hospital_id を持つ
    # ログイン済みユーザーであることのみ確認する。
    jwt_token = credentials.credentials
    user_id = user.get("sub", "")
    _get_hospital_id(user_id, jwt_token)  # hospital_id 未設定ユーザーを弾く

    _remaining()

    # ---- R2 からファイルを取得（Presigned GET、有効期限60秒） ----
    try:
        bucket = get_bucket_name()
        s3 = get_s3_client()
    except Exception:
        logger.exception("R2クライアント初期化失敗 (OCR)")
        raise HTTPException(status_code=500, detail="ストレージ接続エラーが発生しました")

    presigned_url = s3.generate_presigned_url(
        ClientMethod="get_object",
        Params={"Bucket": bucket, "Key": fkey},
        ExpiresIn=60,
    )

    try:
        with urllib.request.urlopen(presigned_url, timeout=10) as resp:
            file_bytes = resp.read()
    except Exception:
        logger.exception("R2からのファイル取得失敗: %s", fkey)
        raise HTTPException(status_code=400, detail="ファイルの取得に失敗しました")

    # ---- サイズチェック（PDF / DOCX 共通） ----
    if len(file_bytes) > _MAX_PDF_SIZE_BYTES:
        mb = len(file_bytes) / 1024 / 1024
        raise HTTPException(
            status_code=400,
            detail=f"ファイルサイズが上限（10MB）を超えています（{mb:.1f}MB）",
        )

    _remaining()

    # ---- ファイル種別ごとのテキスト抽出 ----
    total_pages: Optional[int] = None
    extract_warnings: list[str] = []

    if ext == "docx":
        # DOCX: ローカル抽出（Vision API 不要・高速）
        text, extract_warnings = _extract_docx_text(file_bytes)
        source_type = "docx"

    elif ext == "xlsx":
        # XLSX: openpyxl で全シートをテキスト化（Vision API 不要）
        text, extract_warnings = _extract_xlsx_text(file_bytes)
        source_type = "xlsx"

    else:
        # PDF: pypdfium2 でページ画像化 → Vision OCR
        try:
            png_list, total_pages = _render_pdf_to_png_list(file_bytes)
        except Exception:
            logger.exception("PDF画像化失敗: %s", fkey)
            raise HTTPException(status_code=500, detail="PDF処理でエラーが発生しました")

        if not png_list:
            raise HTTPException(status_code=400, detail="PDFにページが含まれていません")

        if not OPENAI_API_KEY:
            raise HTTPException(
                status_code=500,
                detail="APIキーが未設定です（OPENAI_API_KEY を設定してください）",
            )

        remaining = _remaining()
        text = _call_openai_ocr(png_list, timeout=remaining)

        text = _strip_code_fences(text)
        source_type = "pdf"

    elapsed_ms = int((time.time() - start_time) * 1000)
    stripped = text.strip()

    # ---- AI投入用テキストの正規化（raw は stripped で保持） ----
    _debug_mode = body.mode == "debug"
    normalized, _debug_norm = _normalize_text(stripped, debug=_debug_mode)

    # ---- メタ情報（source_type を追加） ----
    meta = {
        "page_count": total_pages,   # DOCX / XLSX の場合は None（ページ概念なし）
        "char_count": len(stripped),
        "file_key": fkey,
        "elapsed_ms": elapsed_ms,
        "source_type": source_type,  # "pdf" | "docx" | "xlsx"
    }

    # ---- 警告生成（要配慮キーワード検索は normalized を使用） ----
    # extract_warnings: DOCX/XLSX 抽出失敗メッセージがあればそのまま引き継ぐ
    warnings: list[str] = list(extract_warnings)
    if not normalized:
        if not warnings:
            label = "画像OCRでも" if ext == "pdf" else f"{ext.upper()}から"
            warnings.append(
                f"{label}テキストを抽出できませんでした。"
                "内容をご確認の上、送信可否を判断してください。"
            )
    else:
        found = [kw for kw in _SENSITIVE_KEYWORDS if kw in normalized]
        if found:
            warnings.append(
                f"要配慮情報の可能性があります：{', '.join(found)} 等のキーワードが含まれています。"
                "送信前に内容をご確認ください。"
            )

    # ---- 構造化JSON生成（normalized を入力。mode=full のみ実行） ----
    structured = (
        None if body.mode == "text_only" else _structure_referral_text(normalized)
    )

    # ---- アラート生成（normalized を入力。キーワードマッチ方式、断定禁止） ----
    alerts = _generate_alerts(normalized) if normalized else []

    result: dict = {
        "text": stripped,
        "text_normalized": normalized,
        "meta": meta,
        "warnings": warnings,
        "structured": structured,
        "alerts": alerts,
    }
    # debug モードのときのみ debug_normalize を追加（本番レスポンスには含めない）
    if _debug_norm is not None:
        result["debug_normalize"] = _debug_norm
    return result


# ----------------------------
# OCR エンドポイント
# ----------------------------
@app.post("/api/ocr")
def ocr_pdf(
    body: OcrRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """
    POST /api/ocr
    送信前PDFの画像OCR API。AIの判断で送信確定はしない。

    - JWT検証済みユーザーのみ利用可
    - JWT + profiles テーブルで hospital_id 確認（送信前ファイル専用のため documents テーブル照合なし）
    - R2 から Presigned GET でPDF取得（10MB上限）
    - pypdfium2 でページ画像化（最大3ページ）
    - OpenAI Vision API（gpt-4o）で画像OCR
    - OpenAI gpt-4o で構造化JSON生成（OPENAI_API_KEY 未設定時は structured=null）
    - 返却: { text, text_normalized, meta, warnings, structured, alerts }
    """
    return _ocr_impl(body, credentials, user)


@app.post("/ocr")
def ocr_compat(
    body: OcrRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """compat: Vite proxy 経由のローカル開発用（/api/ocr と同じ処理）"""
    return _ocr_impl(body, credentials, user)


# ----------------------------
# 港モデル: Supabase PATCH / POST ヘルパー（user JWT 使用）
# ----------------------------
def _supabase_patch(path: str, data: dict, jwt_token: str) -> list:
    """
    user JWT を使って Supabase REST API を PATCH する（RLS が有効に機能する）。
    Prefer: return=representation → 更新後の行リストを返す。
    0件更新（RLS 拒否 or 条件不一致）の場合は空リストを返す。
    """
    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_ANON_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = json.dumps(data).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        method="PATCH",
        headers={
            "apikey": SUPABASE_ANON_KEY,
            "Authorization": f"Bearer {jwt_token}",
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Prefer": "return=representation",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=5) as resp:
            body_bytes = resp.read()
            return json.loads(body_bytes) if body_bytes else []
    except urllib.error.HTTPError as e:
        logger.error("Supabase user PATCH エラー (%d): path=%s", e.code, path)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")
    except Exception:
        logger.exception("Supabase user PATCH 接続エラー: path=%s", path)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")


def _supabase_post_db(path: str, data: dict, jwt_token: str) -> list:
    """
    user JWT を使って Supabase REST API に POST する（INSERT）。
    RLS INSERT ポリシーが必要。失敗は呼び出し元で try/except すること。
    """
    if not SUPABASE_URL or not SUPABASE_ANON_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_ANON_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = json.dumps(data).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        method="POST",
        headers={
            "apikey": SUPABASE_ANON_KEY,
            "Authorization": f"Bearer {jwt_token}",
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Prefer": "return=representation",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=5) as resp:
            body_bytes = resp.read()
            return json.loads(body_bytes) if body_bytes else []
    except urllib.error.HTTPError as e:
        logger.error("Supabase user POST エラー (%d): path=%s", e.code, path)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")
    except Exception:
        logger.exception("Supabase user POST 接続エラー: path=%s", path)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")


# ----------------------------
# 港モデル: アサイン API
# ----------------------------
class AssignRequest(BaseModel):
    assigned_department: str
    owner_user_id: str                # 担当者の user ID (UUID)
    to_status: Optional[str] = None     # 省略時: 現状維持。"IN_PROGRESS" 推奨


_ASSIGN_VALID_STATUSES = frozenset(
    {"UPLOADED", "IN_PROGRESS", "DOWNLOADED", "ARCHIVED", "CANCELLED"}
)


def _assign_impl(
    doc_id: str,
    body: AssignRequest,
    credentials: HTTPAuthorizationCredentials,
    user: dict,
) -> dict:
    """
    港モデル アサイン処理の共通実装。
    1. JWT → hospital_id 取得
    2. documents GET → to_hospital_id で自院チェック
    3. documents PATCH（owner_user_id / assigned_department / assigned_at / status）
    4. document_logs INSERT（best-effort: 失敗しても 500 にしない）
    """
    jwt_token = credentials.credentials
    user_id = user.get("sub", "")
    hospital_id = _get_hospital_id(user_id, jwt_token)

    # ---- doc_id バリデーション（UUID 形式のみ許可） ----
    doc_id_stripped = doc_id.strip()
    try:
        # UUID 形式チェック（簡易）
        uuid.UUID(doc_id_stripped)
    except ValueError:
        raise HTTPException(status_code=400, detail="無効なドキュメントIDです")

    doc_id_enc = urllib.parse.quote(doc_id_stripped, safe="")

    # ---- 対象ドキュメント取得（RLS: to_hospital_id=自院のみ返る） ----
    rows = _supabase_get(
        f"documents?id=eq.{doc_id_enc}&select=id,status,to_hospital_id,owner_user_id",
        jwt_token,
    )
    if not rows:
        raise HTTPException(
            status_code=404, detail="ドキュメントが見つかりません（権限なし）"
        )

    doc = rows[0]
    if doc.get("to_hospital_id") != hospital_id:
        raise HTTPException(
            status_code=403, detail="自院宛のドキュメントのみアサインできます"
        )

    old_status = doc.get("status", "UPLOADED")
    new_status = body.to_status if body.to_status else old_status
    if new_status not in _ASSIGN_VALID_STATUSES:
        raise HTTPException(status_code=400, detail=f"無効なステータス: {new_status}")

    # ---- assigned_at をサーバー時刻で生成 ----
    assigned_at = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())

    update_data: dict = {
        "owner_user_id": body.owner_user_id,
        "assigned_department": body.assigned_department,
        "assigned_at": assigned_at,
        "status": new_status,
    }

    # ---- PATCH: 自院ドキュメントのみ（to_hospital_id=自院 を URL に追加して二重防御） ----
    hospital_id_enc = urllib.parse.quote(hospital_id, safe="")
    updated = _supabase_patch(
        f"documents?id=eq.{doc_id_enc}&to_hospital_id=eq.{hospital_id_enc}",
        update_data,
        jwt_token,
    )
    if not updated:
        raise HTTPException(
            status_code=403,
            detail="アサインできませんでした（RLS により更新が拒否されました）",
        )

    # ---- document_logs INSERT（best-effort） ----
    try:
        _supabase_post_db(
            "document_logs",
            {
                "document_id": doc_id_stripped,
                "action": "ASSIGN",
                "from_status": old_status,
                "to_status": new_status,
                "changed_by": user_id,
            },
            jwt_token,
        )
    except Exception:
        pass  # best-effort: ログ失敗でも本体処理を継続

    return {
        "ok": True,
        "document_id": doc_id_stripped,
        "assigned_at": assigned_at,
        "owner_user_id": body.owner_user_id,
        "assigned_department": body.assigned_department,
        "status": new_status,
    }


@app.post("/api/documents/{doc_id}/assign")
def assign_document_api(
    doc_id: str,
    body: AssignRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """
    POST /api/documents/{doc_id}/assign
    港モデル: ドキュメントに担当者・部署をアサインする。

    - JWT 必須（自院メンバーのみ）
    - 自院 to_hospital_id チェック（FastAPI + RLS 二重防御）
    - documents: owner_user_id / assigned_department / assigned_at / status を更新
    - document_logs: ASSIGN イベントを記録（best-effort）
    """
    return _assign_impl(doc_id, body, credentials, user)


@app.post("/documents/{doc_id}/assign")
def assign_document_compat(
    doc_id: str,
    body: AssignRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """compat: Vite proxy 経由のローカル開発用（/api/documents/{id}/assign と同じ処理）"""
    return _assign_impl(doc_id, body, credentials, user)


# ===========================================================================
# 変更点（v2.6 CLOUD-FAX-API Webhook統合）:
# 1. POST /api/webhook/cloudfax/inbound: Webhook受信→PDF取得→R2保存→documents INSERT
# 2. fax_inbounds テーブルで冪等性を保証（UNIQUE provider+provider_message_id）
# 3. service_role 使用理由: Webhook は外部システムからの呼び出しのため user JWT が存在しない
#    使用範囲: fax_inbounds / documents の INSERT/PATCH のみ
#
# 変更点（v2.6.1 documents INSERT NOT NULL 対応）:
# 1. to_hospital_id: payload or FAX_DEFAULT_HOSPITAL_ID、どちらも無い場合は 400
# 2. from_hospital_id: MVP では to_hospital_id と同値で入港（NOT NULL を満たす暫定措置）
#    将来: 外部FAX送信元専用の hospital レコードを用意して差し替える
# 3. documents INSERT にメタ情報追加: original_filename / content_type / file_ext / file_size
#
# 変更点（v2.6.2 fax_inbounds に hospital_id を追加）:
# 1. fax_inbounds INSERT に hospital_id = to_hospital_id を追加（病院別受信一覧クエリ対応）
#
# 変更点（v2.7 outbound webhook 追加・PoC モード対応・logger 統一）:
# 1. POST /api/webhook/cloudfax/outbound: FAX送信ステータス通知を受信・保存
#    PDF取得/R2保存/documents INSERT は不要。fax_inbounds に direction=outbound で記録のみ
# 2. _verify_webhook_secret: 共通認証ヘルパー
#    secret 未設定 → PoC モード（起動時警告・処理続行）
#    secret 設定済み → ヘッダー不一致で 401
# 3. print() → logger に統一
# ===========================================================================

# ----------------------------
# CloudFax 環境変数
# ----------------------------
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")
CLOUDFAX_WEBHOOK_SECRET   = os.getenv("CLOUDFAX_WEBHOOK_SECRET", "")   # Webhook認証シークレット
FAX_DEFAULT_HOSPITAL_ID   = os.getenv("FAX_DEFAULT_HOSPITAL_ID", "")   # FAX受信先デフォルト病院ID
# CloudFAX API 認証（実PDF取得に使用）
# CLOUDFAX_BEARER_TOKEN にはトークン本体のみ（"Bearer " プレフィックスは含めない）
CLOUDFAX_API_BASE     = os.getenv("CLOUDFAX_API_BASE", "").rstrip("/")
CLOUDFAX_BEARER_TOKEN = os.getenv("CLOUDFAX_BEARER_TOKEN", "")
CLOUDFAX_API_KEY      = os.getenv("CLOUDFAX_API_KEY", "")
CLOUDFAX_FROM_NUMBER  = os.getenv("CLOUDFAX_FROM_NUMBER", "")  # FAX送信元番号（CloudFAXで払い出した番号）

# ----------------------------
# 起動時セキュリティ警告（F: per-request ではなく起動時1回のみ）
# ----------------------------
# TODO(security): 本番デプロイ前に CLOUDFAX_WEBHOOK_SECRET を必ず設定すること。
#   未設定のまま本番稼働させると、任意のリクエストが Webhook として受理される危険がある。
if not CLOUDFAX_WEBHOOK_SECRET:
    logger.warning(
        "⚠️  [SECURITY / PoC MODE] CLOUDFAX_WEBHOOK_SECRET が未設定です。"
        "CloudFAX Webhook は認証なしで受け付けます。"
        "開発・PoC 環境専用です — 本番デプロイ前に必ず環境変数を設定してください。"
    )


# ----------------------------
# Webhook 共通認証ヘルパー
# ----------------------------
def _verify_webhook_secret(request: Request) -> None:
    """
    CloudFax Webhook の X-CloudFax-Webhook-Secret ヘッダーを検証する。

    - CLOUDFAX_WEBHOOK_SECRET 未設定: PoC モード（起動時に警告済み・処理続行）
    - CLOUDFAX_WEBHOOK_SECRET 設定済み: ヘッダーと一致しない場合 401 を返す

    【適用範囲】
    CloudFAX の inbound webhook 仕様には secret ヘッダが存在しないため、
    inbound エンドポイントではこの関数を呼ばない。
    outbound webhook（送信ステータス通知）にのみ適用する。
    """
    if not CLOUDFAX_WEBHOOK_SECRET:
        # 起動時に警告済み。PoC モードのため認証スキップ。
        return
    secret = request.headers.get("X-CloudFax-Webhook-Secret", "")
    if secret != CLOUDFAX_WEBHOOK_SECRET:
        client_host = request.client.host if request.client else "unknown"
        logger.warning("[cloudfax] Webhook シークレット不一致: remote=%s", client_host)
        raise HTTPException(status_code=401, detail="Webhook シークレットが無効です")


# ----------------------------
# service_role 用 Supabase ヘルパー（Webhook処理専用）
# service_role は fax_inbounds / documents の INSERT/PATCH のみに限定して使用する
# ----------------------------
def _supabase_service_post(path: str, data: dict, prefer: str = "return=representation") -> list:
    """service_role で Supabase REST API に POST する（Webhook処理専用）"""
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = json.dumps(data).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        method="POST",
        headers={
            "apikey":        SUPABASE_SERVICE_ROLE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
            "Content-Type":  "application/json",
            "Accept":        "application/json",
            "Prefer":        prefer,
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            body_bytes = resp.read()
            return json.loads(body_bytes) if body_bytes else []
    except urllib.error.HTTPError as e:
        body_text = e.read().decode(errors="replace")
        logger.error("Supabase service POST エラー (%d): %s", e.code, body_text)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")
    except Exception:
        logger.exception("Supabase service POST 接続エラー")
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")


def _supabase_service_patch(path: str, data: dict) -> list:
    """service_role で Supabase REST API を PATCH する（Webhook処理専用）"""
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    payload = json.dumps(data).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        method="PATCH",
        headers={
            "apikey":        SUPABASE_SERVICE_ROLE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
            "Content-Type":  "application/json",
            "Accept":        "application/json",
            "Prefer":        "return=representation",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            body_bytes = resp.read()
            return json.loads(body_bytes) if body_bytes else []
    except urllib.error.HTTPError as e:
        logger.error("Supabase service PATCH エラー (%d)", e.code)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")
    except Exception:
        logger.exception("Supabase service PATCH 接続エラー")
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")


def _supabase_service_get(path: str) -> list:
    """service_role で Supabase REST API を GET する（Webhook処理専用）"""
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise HTTPException(
            status_code=500,
            detail="SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY が未設定です",
        )
    url = f"{SUPABASE_URL.rstrip('/')}/rest/v1/{path}"
    req = urllib.request.Request(
        url,
        method="GET",
        headers={
            "apikey":        SUPABASE_SERVICE_ROLE_KEY,
            "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
            "Accept":        "application/json",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            body_bytes = resp.read()
            return json.loads(body_bytes) if body_bytes else []
    except urllib.error.HTTPError as e:
        body_text = e.read().decode(errors="replace")
        logger.error("Supabase service GET エラー (%d): %s", e.code, body_text)
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")
    except Exception:
        logger.exception("Supabase service GET 接続エラー")
        raise HTTPException(status_code=502, detail="データベース操作でエラーが発生しました")


# ----------------------------
# R2 直接アップロードヘルパー（Webhook→PDF保存専用）
# ----------------------------
def _r2_put_object(file_key: str, data: bytes, content_type: str = "application/pdf") -> None:
    """R2 に直接 PUT する（Presigned URL 不使用。Webhook処理でのみ使用）"""
    try:
        bucket = get_bucket_name()
        s3 = get_s3_client()
    except Exception:
        logger.exception("R2クライアント初期化失敗 (put_object)")
        raise HTTPException(status_code=500, detail="ストレージ接続エラーが発生しました")
    s3.put_object(Bucket=bucket, Key=file_key, Body=data, ContentType=content_type)


# ----------------------------
# CloudFax API ヘルパー（実API呼び出し専用）
# ----------------------------

def _cloudfax_auth_headers(accept: str = "application/json") -> dict:
    """
    CloudFax API 認証ヘッダを生成する。
    環境変数 CLOUDFAX_API_BASE / CLOUDFAX_BEARER_TOKEN / CLOUDFAX_API_KEY が
    未設定の場合は RuntimeError を raise する。
    ※ CLOUDFAX_BEARER_TOKEN にはトークン本体のみ（"Bearer " プレフィックスは含めない）
    """
    if not CLOUDFAX_API_BASE or not CLOUDFAX_BEARER_TOKEN or not CLOUDFAX_API_KEY:
        raise RuntimeError(
            "CloudFAX API 設定が不足しています "
            "(CLOUDFAX_API_BASE / CLOUDFAX_BEARER_TOKEN / CLOUDFAX_API_KEY)"
        )
    return {
        "Accept":        accept,
        "Authorization": f"Bearer {CLOUDFAX_BEARER_TOKEN}",
        "x-api-key":     CLOUDFAX_API_KEY,
    }


def _cloudfax_fetch_status(transmission_id: str) -> dict:
    """
    GET /v1/Faxes/{transmission_id} で FAX ステータス JSON を取得する。
    media_url を含む CloudFAX のレスポンス dict を返す。
    """
    url = f"{CLOUDFAX_API_BASE}/Faxes/{urllib.parse.quote(transmission_id, safe='')}"
    logger.info("[cloudfax] ステータス取得: transmission_id=%s", transmission_id)
    req = urllib.request.Request(url, headers=_cloudfax_auth_headers("application/json"))
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            body = json.loads(resp.read().decode())
            logger.debug("[cloudfax] ステータスレスポンス: status=%s", body.get("status"))
            return body
    except urllib.error.HTTPError as e:
        body_text = e.read().decode(errors="replace")
        logger.error("[cloudfax] ステータス取得 HTTP エラー (%d): %s", e.code, body_text)
        raise RuntimeError(f"CloudFAX ステータス取得失敗 (HTTP {e.code})")
    except RuntimeError:
        raise
    except Exception as e:
        logger.exception("[cloudfax] ステータス取得 接続エラー")
        raise RuntimeError(f"CloudFAX ステータス取得 接続エラー: {e}")


def _cloudfax_fetch_media(media_url: str) -> bytes:
    """
    media_url から PDF bytes を取得する。
    Accept: application/pdf,application/octet-stream,application/json
    レスポンスが PDF ではなく JSON だった場合はログを出して RuntimeError を raise する。
    """
    safe_url = media_url.split("?")[0]
    logger.info("[cloudfax] PDF取得開始: media_url=%s", safe_url)
    req = urllib.request.Request(
        media_url,
        headers=_cloudfax_auth_headers(
            "application/pdf,application/octet-stream,application/json"
        ),
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            content_type = resp.headers.get("Content-Type", "")
            data = resp.read()
            # Content-Type が JSON かつ %PDF ヘッダが無い場合は仕様齟齬として失敗扱い
            if not data.startswith(b"%PDF") and "json" in content_type.lower():
                logger.error(
                    "[cloudfax] media_url が JSON を返しました (Content-Type=%s 先頭=%s)",
                    content_type, data[:120],
                )
                raise RuntimeError(
                    f"media_url が PDF ではなく JSON を返しました "
                    f"(Content-Type={content_type!r})"
                )
            logger.info(
                "[cloudfax] PDF取得完了: size=%d bytes, Content-Type=%s",
                len(data), content_type,
            )
            return data
    except urllib.error.HTTPError as e:
        body_text = e.read().decode(errors="replace")
        logger.error("[cloudfax] PDF取得 HTTP エラー (%d): %s", e.code, body_text)
        raise RuntimeError(f"CloudFAX PDF取得失敗 (HTTP {e.code})")
    except RuntimeError:
        raise
    except Exception as e:
        logger.exception("[cloudfax] PDF取得 接続エラー")
        raise RuntimeError(f"CloudFAX PDF取得 接続エラー: {e}")


# ----------------------------
# A. PDF 妥当性確認ヘルパー
# ----------------------------
def _validate_pdf_bytes(data: bytes, source: str = "") -> None:
    """
    PDF バイト列が有効かを確認する（R2 保存前チェック）。
    - %PDF ヘッダ確認
    - pypdfium2 で開けるか確認
    - ページ数 >= 1 確認
    - 先頭ページへのアクセス確認（ページオブジェクトが壊れていないか）
    壊れていた場合は ValueError を raise する。
    呼び出し元で ValueError を捕捉して error_stage=PDF_VALIDATE として記録すること。
    """
    if not data.startswith(b"%PDF"):
        raise ValueError(f"不正な PDF: %PDF ヘッダがありません (source={source!r})")
    try:
        doc = pdfium.PdfDocument(data)
        page_count = len(doc)
        if page_count < 1:
            raise ValueError(f"PDF にページがありません (source={source!r})")
        _ = doc[0]  # 先頭ページオブジェクトへのアクセス確認（重いレンダリングは不要）
        logger.debug("[pdf_validate] OK: pages=%d source=%s", page_count, source)
    except ValueError:
        raise  # 上で raise した ValueError はそのまま伝播
    except Exception as e:
        raise ValueError(f"pypdfium2 で PDF を開けません: {e} (source={source!r})")


# ----------------------------
# CloudFax PDF 取得（実API実装）
# ----------------------------
async def fetch_pdf_from_cloudfax(
    provider_message_id: str,
    payload_raw: dict | None = None,
) -> bytes:
    """
    CloudFax API から PDF bytes を取得する。

    取得優先順:
      1. payload_raw["media_url"] が存在すればそのまま GET → PDF bytes
      2. なければ payload_raw["transmission_id"] または provider_message_id を使って
         GET /v1/Faxes/{TransmissionId} を叩き、レスポンスの media_url を取得してから GET

    取得後は呼び出し元 _cloudfax_inbound_impl() で _validate_pdf_bytes() が実行される。
    """
    # 1. payload から media_url を優先取得
    media_url: str = str((payload_raw or {}).get("media_url") or "").strip()

    if not media_url:
        # 2. transmission_id → GET /v1/Faxes/{id} → media_url を取得（フォールバック）
        transmission_id = str(
            (payload_raw or {}).get("transmission_id") or provider_message_id
        ).strip()
        if not transmission_id:
            raise RuntimeError("transmission_id / provider_message_id が取得できません")

        status_json = _cloudfax_fetch_status(transmission_id)
        media_url   = str(status_json.get("media_url") or "").strip()
        if not media_url:
            raise RuntimeError(
                f"CloudFAX ステータスレスポンスに media_url がありません "
                f"(transmission_id={transmission_id!r})"
            )
        logger.info(
            "[cloudfax] フォールバック: ステータスAPIから media_url を取得 "
            "(transmission_id=%s)",
            transmission_id,
        )

    return _cloudfax_fetch_media(media_url)


# ----------------------------
# CloudFax Webhook Pydantic モデル（最小限）
# ----------------------------
class CloudFaxPayload(BaseModel):
    """CloudFax Inbound Webhook payload の既知フィールド定義"""
    id: Optional[str] = None   # CloudFax が送ってくる場合
    fax_id: Optional[str] = None   # 別プロバイダが fax_id を使う場合
    to_hospital_id: Optional[str] = None   # 受信先病院ID（省略時は FAX_DEFAULT_HOSPITAL_ID）
    model_config = {"extra": "allow"}   # 未知フィールドは raw として保存


# ----------------------------
# C. エラーステージ定数（fax_inbounds.error_stage に保存する値）
# ----------------------------
_STAGE_VALIDATION      = "VALIDATION"
_STAGE_PDF_FETCH       = "PDF_FETCH"
_STAGE_PDF_VALIDATE    = "PDF_VALIDATE"
_STAGE_R2_UPLOAD       = "R2_UPLOAD"
_STAGE_DOCUMENT_INSERT = "DOCUMENT_INSERT"
_STAGE_STATUS_UPDATE   = "STATUS_UPDATE"


# ----------------------------
# CloudFax Inbound Webhook 処理実装
# ----------------------------
async def _cloudfax_inbound_impl(payload_raw: dict, background_tasks: BackgroundTasks) -> dict:
    """
    CloudFax Inbound Webhook の共通処理。

    フロー:
      1. provider_message_id 取得（id / fax_id の優先順）
      2. fax_inbounds を GET で検索し、既存行の有無・ステータスで分岐する
         - 既存行なし        → 新規 INSERT して処理続行
         - status=FAILED     → PATCH でリセット（RECEIVED / error=NULL）して再処理続行（retry）
         - status=DOC_CREATED / RECEIVED / その他 → 冪等スキップ（即返却）
      3. PDF 取得（media_url 優先 / なければ GET /v1/Faxes/{id} でフォールバック）
         → PDF 妥当性確認（%PDF ヘッダ + pypdfium2 open チェック）
      4. R2 保存
      5. documents INSERT（status=ARRIVED, owner_user_id=NULL）
      6. fax_inbounds.status を DOC_CREATED に更新
      例外時: fax_inbounds.status を FAILED + error_stage に更新
    """
    # provider_message_id を id / fax_id / transmission_id から取得（優先順）
    provider_message_id = str(
        payload_raw.get("id")
        or payload_raw.get("fax_id")
        or payload_raw.get("transmission_id")
        or ""
    ).strip()
    if not provider_message_id:
        raise HTTPException(status_code=400, detail="payload に id / fax_id / transmission_id が必要です")

    # to_hospital_id: payload → 環境変数 → 400 の優先順
    to_hospital_id = (payload_raw.get("to_hospital_id") or FAX_DEFAULT_HOSPITAL_ID or "").strip()
    if not to_hospital_id:
        raise HTTPException(
            status_code=400,
            detail="to_hospital_id required: payload に to_hospital_id を含めるか FAX_DEFAULT_HOSPITAL_ID を設定してください",
        )

    # ---- 既存行チェック（FAILED 再処理対応）----
    # GET で既存行を先に確認し、ステータスに応じて分岐する。
    # FAILED 行は PATCH でリセットして再処理続行、それ以外は冪等返却。
    msg_enc  = urllib.parse.quote(provider_message_id, safe="")
    existing = _supabase_service_get(
        f"fax_inbounds?provider=eq.cloudfax&provider_message_id=eq.{msg_enc}&select=id,status"
    )

    if existing:
        existing_row    = existing[0]
        existing_status = existing_row.get("status", "")
        fax_inbound_id  = existing_row["id"]

        if existing_status == "FAILED":
            # FAILED → error をリセットして再処理続行
            fax_enc = urllib.parse.quote(fax_inbound_id, safe="")
            _supabase_service_patch(
                f"fax_inbounds?id=eq.{fax_enc}",
                {"status": "RECEIVED", "error": None, "error_stage": None},
            )
            logger.info(
                "[cloudfax] FAILED再処理 (retry): fax_inbound_id=%s, msg_id=%s",
                fax_inbound_id, provider_message_id,
            )
        else:
            # DOC_CREATED / RECEIVED / その他の未知ステータス → 冪等返却（安全側）
            logger.info(
                "[cloudfax] 冪等(status=%s): provider_message_id=%s は処理済み",
                existing_status, provider_message_id,
            )
            return {"ok": True, "idempotent": True, "provider_message_id": provider_message_id}

    else:
        # 既存行なし → 新規 INSERT
        inserted = _supabase_service_post(
            "fax_inbounds",
            {
                "provider":            "cloudfax",
                "provider_message_id": provider_message_id,
                "direction":           "inbound",
                "status":              "RECEIVED",
                "hospital_id":         to_hospital_id,   # 病院別受信一覧のためのインデックスキー
                "raw":                 payload_raw,
            },
        )
        if not inserted:
            raise RuntimeError("fax_inbounds INSERT に失敗しました（レスポンスが空）")
        fax_inbound_id = inserted[0]["id"]
        logger.info("[cloudfax] 新規受信: fax_inbound_id=%s, msg_id=%s", fax_inbound_id, provider_message_id)

    error_stage = _STAGE_PDF_FETCH  # C: 失敗時にどの段階か追跡する
    try:
        # ---- PDF 取得（payload_raw を渡し media_url を優先利用）----
        pdf_bytes = await fetch_pdf_from_cloudfax(provider_message_id, payload_raw)

        # ---- A. PDF 妥当性確認（R2 保存前に壊れた PDF を検出する）----
        error_stage = _STAGE_PDF_VALIDATE
        _validate_pdf_bytes(pdf_bytes, source=provider_message_id)

        # ---- R2 保存 ----
        error_stage = _STAGE_R2_UPLOAD
        file_key = f"documents/{uuid.uuid4()}.pdf"
        _r2_put_object(file_key, pdf_bytes)
        logger.info("[cloudfax] R2 保存完了: file_key=%s", file_key)

        # ---- documents INSERT ----
        # status=ARRIVED: 港モデルの「未担当BOX」に入港する
        # owner_user_id=NULL: 担当者未割り当て（アサイン機能で後から設定）
        # from_hospital_id=to_hospital_id: FAX送信元病院は不明のため受信先と同値（NOT NULL 暫定措置）
        #   将来: 外部FAX送信元専用の hospital レコードを作成し、そちらの hospital_id を設定する
        error_stage = _STAGE_DOCUMENT_INSERT
        doc_rows = _supabase_service_post(
            "documents",
            {
                "file_key":          file_key,
                "status":            "ARRIVED",
                "owner_user_id":     None,
                "from_hospital_id":  to_hospital_id,   # 暫定: 送信元不明のため受信先で代替
                "to_hospital_id":    to_hospital_id,
                "original_filename": f"fax_{provider_message_id}.pdf",
                "content_type":      "application/pdf",
                "file_ext":          "pdf",
                "file_size":         len(pdf_bytes),
                "from_fax_number":   payload_raw.get("from") or None,
                "to_fax_number":     payload_raw.get("to")   or None,
                "source":            "fax",
                "ocr_status":        "PENDING",
            },
        )
        if not doc_rows:
            raise RuntimeError("documents INSERT に失敗しました（レスポンスが空）")
        doc_id = doc_rows[0]["id"]
        logger.info("[cloudfax] documents INSERT 完了: doc_id=%s", doc_id)

        # ---- バックグラウンドOCR（best-effort: 失敗してもWebhook応答は成功） ----
        if OPENAI_API_KEY:
            background_tasks.add_task(_analyze_document_for_fax, doc_id, file_key)
            logger.info("[cloudfax] OCRバックグラウンドタスク登録: doc_id=%s", doc_id)
        else:
            logger.warning("[cloudfax] OPENAI_API_KEY 未設定のためOCRスキップ: doc_id=%s", doc_id)

        # ---- fax_inbounds を DOC_CREATED に更新 ----
        error_stage = _STAGE_STATUS_UPDATE
        fax_enc = urllib.parse.quote(fax_inbound_id, safe="")
        _supabase_service_patch(
            f"fax_inbounds?id=eq.{fax_enc}",
            {"status": "DOC_CREATED", "document_id": doc_id, "file_key": file_key},
        )

        return {
            "ok":             True,
            "fax_inbound_id": fax_inbound_id,
            "document_id":    doc_id,
            "file_key":       file_key,
        }

    except HTTPException:
        raise  # FastAPI エラーはそのまま伝播
    except Exception as e:
        # ---- エラー時: fax_inbounds を FAILED + error_stage に更新 ----
        logger.error(
            "[cloudfax] エラー stage=%s fax_inbound_id=%s: %s",
            error_stage, fax_inbound_id, e,
        )
        try:
            fax_enc = urllib.parse.quote(fax_inbound_id, safe="")
            _supabase_service_patch(
                f"fax_inbounds?id=eq.{fax_enc}",
                {"status": "FAILED", "error": str(e)[:500], "error_stage": error_stage},
            )
        except Exception:
            logger.exception("[cloudfax] FAILED ステータス更新にも失敗")
        logger.exception("CloudFax Webhook 処理エラー")
        raise HTTPException(status_code=500, detail="Webhook 処理でエラーが発生しました")


# ----------------------------
# FAX受信文書 バックグラウンドOCR
# ----------------------------
# 紹介状判定キーワード（OCRテキストに含まれれば "紹介状" と分類）
_REFERRAL_KEYWORDS = [
    "紹介状", "診療情報提供書", "診療情報提供", "ご紹介", "御紹介",
    "紹介先", "紹介元", "かかりつけ", "専門診療科",
]

_MAX_FAX_OCR_SECS = 90  # FAX OCRのタイムアウト（秒）


def _analyze_document_for_fax(document_id: str, file_key: str) -> None:
    """
    FAX受信PDFに対してOCR + document_type分類を実行し、documentsを更新する。
    - BackgroundTasks から呼ばれる（同期関数）
    - 失敗しても documents 登録は影響しない（best-effort）
    - document_type: "紹介状" | "不明"
    """
    logger.info("[fax-ocr] 開始: document_id=%s file_key=%s", document_id, file_key)
    try:
        # ---- R2 からPDF取得 ----
        try:
            bucket = get_bucket_name()
            s3 = get_s3_client()
        except Exception:
            logger.exception("[fax-ocr] R2クライアント初期化失敗")
            return

        presigned_url = s3.generate_presigned_url(
            ClientMethod="get_object",
            Params={"Bucket": bucket, "Key": file_key},
            ExpiresIn=120,
        )
        try:
            with urllib.request.urlopen(presigned_url, timeout=15) as resp:
                pdf_bytes = resp.read()
        except Exception:
            logger.exception("[fax-ocr] R2からのPDF取得失敗: %s", file_key)
            return

        if len(pdf_bytes) > _MAX_PDF_SIZE_BYTES:
            logger.warning("[fax-ocr] PDFサイズ超過 (%d bytes), スキップ", len(pdf_bytes))
            return

        # ---- PDF → PNG → OCR ----
        try:
            png_list, _ = _render_pdf_to_png_list(pdf_bytes)
        except Exception:
            logger.exception("[fax-ocr] PDF画像化失敗: %s", file_key)
            _supabase_service_patch(
                f"documents?id=eq.{urllib.parse.quote(document_id, safe='')}",
                {"ocr_status": "FAILED"},
            )
            return

        logger.info("[fax-ocr] OCR開始: pages=%d document_id=%s", len(png_list), document_id)
        _supabase_service_patch(
            f"documents?id=eq.{urllib.parse.quote(document_id, safe='')}",
            {"ocr_status": "RUNNING"},
        )
        try:
            raw_text = _call_openai_ocr(png_list, timeout=_MAX_FAX_OCR_SECS)
        except Exception:
            logger.exception("[fax-ocr] OpenAI OCR失敗: %s", file_key)
            _supabase_service_patch(
                f"documents?id=eq.{urllib.parse.quote(document_id, safe='')}",
                {"ocr_status": "FAILED"},
            )
            return

        logger.info("[fax-ocr] OCR完了: chars=%d document_id=%s", len(raw_text), document_id)
        normalized, _ = _normalize_text(raw_text)

        # ---- document_type 分類 ----
        doc_type = "不明"
        for kw in _REFERRAL_KEYWORDS:
            if kw in normalized or kw in raw_text:
                doc_type = "紹介状"
                break

        # ---- structured_json 生成（失敗時は None のまま） ----
        structured = _structure_referral_text(normalized)
        if structured:
            logger.info("[fax-ocr] 構造化JSON生成完了: document_id=%s", document_id)
        else:
            logger.warning("[fax-ocr] 構造化JSON生成スキップ/失敗: document_id=%s", document_id)

        # ---- documents を更新 ----
        patch_data: dict = {
            "ocr_text":      raw_text,
            "ocr_status":    "DONE",
            "document_type": doc_type,
        }
        if structured:
            patch_data["structured_json"] = structured

        _supabase_service_patch(
            f"documents?id=eq.{urllib.parse.quote(document_id, safe='')}",
            patch_data,
        )
        logger.info(
            "[fax-ocr] 完了: document_id=%s document_type=%s",
            document_id, doc_type,
        )

    except Exception:
        logger.exception("[fax-ocr] 予期しないエラー: document_id=%s", document_id)


# ----------------------------
# CloudFax Webhook エンドポイント
# ----------------------------
@app.post("/api/webhook/cloudfax/inbound")
async def cloudfax_inbound_api(request: Request, background_tasks: BackgroundTasks):
    """
    POST /api/webhook/cloudfax/inbound
    CloudFax からの Inbound FAX Webhook を受信する。

    - 認証: CloudFAX inbound webhook は secret ヘッダを送付しない仕様のため、
            X-CloudFax-Webhook-Secret による検証は行わない
    - 冪等性: fax_inbounds の UNIQUE(provider, provider_message_id) で保証
    - service_role 使用: user JWT が存在しない外部Webhook処理のため（最小範囲）
    """
    try:
        payload_raw: dict = await request.json()
    except Exception:
        logger.error("[cloudfax/inbound] JSON パース失敗")
        raise HTTPException(status_code=400, detail="JSON パースに失敗しました")

    # 切り分け用ログ: payload キー一覧と安全なフィールドのみ出力（秘密情報は含めない）
    logger.info("[cloudfax/inbound] payload keys=%s", list(payload_raw.keys()))
    logger.info(
        "[cloudfax/inbound] transmission_id=%s status=%s direction=%s media_url_present=%s",
        payload_raw.get("transmission_id"),
        payload_raw.get("status"),
        payload_raw.get("direction"),
        bool(payload_raw.get("media_url")),
    )

    # E: Pydantic モデルでフィールド型確認（extra="allow" で未知フィールドも通過）
    try:
        CloudFaxPayload.model_validate(payload_raw)
    except Exception as e:
        logger.error("[cloudfax/inbound] payload バリデーション失敗: %s", e)
        raise HTTPException(status_code=400, detail="payload のバリデーションに失敗しました")

    return await _cloudfax_inbound_impl(payload_raw, background_tasks)


@app.post("/webhook/cloudfax/inbound")
async def cloudfax_inbound_compat(request: Request, background_tasks: BackgroundTasks):
    """compat: Vite proxy 経由のローカル開発用（/api/webhook/cloudfax/inbound と同じ処理）"""
    try:
        payload_raw: dict = await request.json()
    except Exception:
        logger.error("[cloudfax/inbound] JSON パース失敗")
        raise HTTPException(status_code=400, detail="JSON パースに失敗しました")

    try:
        CloudFaxPayload.model_validate(payload_raw)
    except Exception as e:
        logger.error("[cloudfax/inbound] payload バリデーション失敗: %s", e)
        raise HTTPException(status_code=400, detail="payload のバリデーションに失敗しました")

    return await _cloudfax_inbound_impl(payload_raw, background_tasks)


# ----------------------------
# CloudFax Outbound Webhook モデル・実装
# ----------------------------
class CloudFaxOutboundPayload(BaseModel):
    """
    CloudFax Outbound Webhook payload（FAX送信ステータス通知）の既知フィールド定義。
    # TODO(cloudfax-spec): CloudFAX の実際の outbound payload 仕様が確定したら
    #   フィールド名・status 値を合わせて修正すること。
    """
    id: Optional[str] = None  # CloudFax が送ってくる場合
    fax_id: Optional[str] = None   # 別プロバイダが fax_id を使う場合
    status: Optional[str] = None   # SENT / FAILED / DELIVERING など（仕様未確定）
    hospital_id: Optional[str] = None   # 送信元病院ID（省略時は FAX_DEFAULT_HOSPITAL_ID）
    model_config = {"extra": "allow"}  # 未知フィールドは raw として保存


async def _cloudfax_outbound_impl(payload_raw: dict) -> dict:
    """
    CloudFax Outbound Webhook（FAX送信ステータス通知）の共通処理。

    【設計】B: outbound は同一 FAX に対して複数ステータス通知が来る
    （例: QUEUED → SENDING → SENT / FAILED）。
    fax_inbounds の (provider, provider_message_id) 単位では後続通知が潰れるため、
    fax_webhook_events テーブルに (provider, provider_message_id, event_status) 単位で記録する。

    フロー:
      1. provider_message_id 取得（id / fax_id の優先順）
      2. event_status を payload.status から取得（未定義なら UNKNOWN）
      3. fax_webhook_events に INSERT（UNIQUE: provider + provider_message_id + event_status）
         → 同一 FAX の別ステータスは別行として記録（status 遷移履歴が追える）
         → 同一 FAX + 同一 status の重複通知は冪等（200 OK）

    TODO(v2.8): 送信済みドキュメントのステータス連携が必要な場合は、
      provider_message_id でドキュメントを特定して status を更新する処理を追加すること。
      そのためには documents テーブルに provider_message_id カラムを追加し、
      FAX 送信時に保存する仕組みが必要。
    """
    provider_message_id = str(
        payload_raw.get("id") or payload_raw.get("fax_id") or ""
    ).strip()
    if not provider_message_id:
        logger.error(
            "[cloudfax/outbound] payload に id または fax_id がありません: keys=%s",
            list(payload_raw.keys()),
        )
        raise HTTPException(status_code=400, detail="payload に id または fax_id が必要です")

    # TODO(cloudfax-spec): status フィールド名は CloudFAX の実仕様書で確認すること
    event_status = str(payload_raw.get("status") or "UNKNOWN").strip().upper()

    # hospital_id は outbound では任意（省略時は FAX_DEFAULT_HOSPITAL_ID、それも無ければ NULL）
    hospital_id: Optional[str] = (
        payload_raw.get("hospital_id") or FAX_DEFAULT_HOSPITAL_ID or ""
    ).strip() or None

    # ---- fax_webhook_events INSERT（冪等: provider + provider_message_id + event_status 単位）----
    # 同一 FAX への複数ステータス通知（QUEUED/SENDING/SENT 等）を個別に記録する。
    inserted = _supabase_service_post(
        "fax_webhook_events?on_conflict=provider,provider_message_id,event_status",
        {
            "provider":            "cloudfax",
            "provider_message_id": provider_message_id,
            "direction":           "outbound",
            "event_status":        event_status,
            "hospital_id":         hospital_id,
            "raw":                 payload_raw,
        },
        prefer="resolution=ignore-duplicates,return=representation",
    )

    if not inserted:
        logger.info(
            "[cloudfax/outbound] 冪等: msg_id=%s status=%s は既録",
            provider_message_id, event_status,
        )
        return {
            "ok":                  True,
            "idempotent":          True,
            "provider_message_id": provider_message_id,
            "event_status":        event_status,
        }

    event_id = inserted[0]["id"]
    logger.info(
        "[cloudfax/outbound] 新規イベント: event_id=%s msg_id=%s status=%s",
        event_id, provider_message_id, event_status,
    )

    return {
        "ok":                  True,
        "fax_event_id":        event_id,
        "provider_message_id": provider_message_id,
        "event_status":        event_status,
    }


@app.post("/api/webhook/cloudfax/outbound")
async def cloudfax_outbound_api(request: Request):
    """
    POST /api/webhook/cloudfax/outbound
    CloudFax からの FAX送信ステータス通知 Webhook を受信する。

    - 認証: X-CloudFax-Webhook-Secret ヘッダー（_verify_webhook_secret 参照）
    - 冪等性: fax_webhook_events の UNIQUE(provider, provider_message_id, event_status) で保証
      → 同一 FAX への複数 status 通知（QUEUED/SENDING/SENT 等）は個別に記録される
    - PDF取得・R2保存・documents INSERT は行わない（ステータスイベント記録のみ）
    """
    _verify_webhook_secret(request)

    try:
        payload_raw: dict = await request.json()
    except Exception:
        logger.error("[cloudfax/outbound] JSON パース失敗")
        raise HTTPException(status_code=400, detail="JSON パースに失敗しました")

    # E: Pydantic モデルでフィールド型確認（extra="allow" で未知フィールドも通過）
    try:
        CloudFaxOutboundPayload.model_validate(payload_raw)
    except Exception as e:
        logger.error("[cloudfax/outbound] payload バリデーション失敗: %s", e)
        raise HTTPException(status_code=400, detail="payload のバリデーションに失敗しました")

    return await _cloudfax_outbound_impl(payload_raw)


# ===========================================================================
# FAX送信エンドポイント（v2.8 追加）
# POST /api/send-fax
# 認証: Supabase JWT 必須（_bearer / verify_jwt）
# 処理: R2上の既アップロード済みファイルを CloudFAX API 経由で FAX 送信する
# ===========================================================================

class SendFaxRequest(BaseModel):
    file_key:          str
    contact_id:        str               # contacts.id（サーバ側で fax_number を取得するために使用）
    fax_number:        Optional[str] = None  # Deprecated: サーバ側で contacts から取得するため無視される
    comment:           Optional[str] = None
    original_filename: Optional[str] = None


async def _send_fax_impl(
    req: SendFaxRequest,
    hospital_id: str,
    user_id: str,
    jwt_token: str,
    fax_number: str,   # サーバ側で contacts から解決済みの FAX番号
) -> dict:
    """
    CloudFAX API で FAX 送信する共通処理。

    フロー:
    1. R2 の presigned GET URL を MediaUrl として生成（PDF取得は CloudFAX 側に委譲）
    2. CloudFAX POST /v1/Faxes で送信依頼（application/json）
    3. documents テーブルに source="fax_outbound" で記録
    4. document_events に FAX_SEND を記録（best-effort）
    """
    # 1. CloudFAX 設定チェック（presigned URL 生成前に失敗させる）
    if not CLOUDFAX_API_BASE or not CLOUDFAX_BEARER_TOKEN or not CLOUDFAX_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="CloudFAX API が未設定です (CLOUDFAX_API_BASE / CLOUDFAX_BEARER_TOKEN / CLOUDFAX_API_KEY)",
        )
    if not CLOUDFAX_FROM_NUMBER:
        raise HTTPException(
            status_code=503,
            detail="FAX送信元番号が未設定です (CLOUDFAX_FROM_NUMBER)",
        )

    # 2. R2 の presigned GET URL を生成（CloudFAX が PDF を取得するための MediaUrl）
    #    有効期限: 600s（CloudFAX がダウンロードしに来るまでの余裕を持たせる）
    try:
        bucket = get_bucket_name()
        s3     = get_s3_client()
    except Exception:
        logger.exception("[send-fax] R2クライアント初期化失敗")
        raise HTTPException(status_code=500, detail="ストレージ接続エラー")

    media_url = s3.generate_presigned_url(
        ClientMethod="get_object",
        Params={"Bucket": bucket, "Key": req.file_key},
        ExpiresIn=600,
    )

    # 3. CloudFAX API で送信依頼（POST /v1/Faxes, application/json）
    send_url  = f"{CLOUDFAX_API_BASE}/Faxes"
    body_dict = {"From": CLOUDFAX_FROM_NUMBER, "To": fax_number, "MediaUrl": media_url}
    body_bytes = json.dumps(body_dict).encode()

    logger.info(
        "[send-fax] CloudFAX送信直前: url=%s From=%s To=%s MediaUrl(先頭80文字)=%s",
        send_url, CLOUDFAX_FROM_NUMBER, fax_number, media_url[:80],
    )

    send_req = urllib.request.Request(
        send_url,
        data=body_bytes,
        method="POST",
        headers={
            "Accept":        "application/json",
            "Authorization": f"Bearer {CLOUDFAX_BEARER_TOKEN}",
            "Content-Type":  "application/json",
            "x-api-key":     CLOUDFAX_API_KEY,
        },
    )
    try:
        with urllib.request.urlopen(send_req, timeout=60) as resp:
            send_result = json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read().decode(errors="replace")
        except Exception:
            err_body = "(body読み取り失敗)"
        logger.error(
            "[send-fax] CloudFAX HTTPError: status=%s reason=%s body=%s",
            e.code, e.reason, err_body,
        )
        raise HTTPException(status_code=502, detail=f"FAX送信APIエラー: {e.code} {err_body}")
    except Exception as e:
        logger.exception("[send-fax] CloudFAX送信 予期せぬ例外")
        raise HTTPException(status_code=502, detail=f"FAX送信APIエラー: {e}")

    transmission_id = send_result.get("transmission_id") or send_result.get("id") or ""
    logger.info("[send-fax] CloudFAX送信依頼完了: transmission_id=%s to=%s", transmission_id, fax_number)
    # TODO(transmission_id): documents に transmission_id カラムを追加し、ここで保存すること。
    #   そうすれば outbound webhook と documents を紐付けられる。

    # 3. documents テーブルに記録（source="fax_outbound"）
    # service_role 使用理由: JWT ユーザーのスコープ外テーブル行を書くため（RLS バイパス）
    doc_rows = _supabase_service_post(
        "documents",
        {
            "from_hospital_id":  hospital_id,
            "to_hospital_id":    hospital_id,   # FAX相手は hospitals 外のため自院IDで代替
            "to_fax_number":     fax_number,
            # TODO(doc_insert_failure): CloudFAX送信成功後にここが失敗した場合、
            #   FAXは送信済みだが documents レコードが存在しない状態になる。
            #   MVP では best-effort（失敗をログして続行）とするが、
            #   将来的には transmission_id を使って冪等 upsert に昇格させること。
            "file_key":          req.file_key,
            "original_filename": req.original_filename,
            "comment":           req.comment,
            "status":            "UPLOADED",
            "source":            "fax_outbound",
        },
    )
    doc_id = doc_rows[0]["id"] if doc_rows else ""

    # 4. 監査ログ（best-effort）
    if doc_id:
        try:
            _supabase_service_post(
                "document_events",
                {
                    "document_id": doc_id,
                    "user_id":     user_id,
                    "event_type":  "FAX_SEND",
                    "hospital_id": hospital_id,
                },
            )
        except Exception:
            pass

    return {
        "ok":              True,
        "transmission_id": transmission_id,
        "document_id":     doc_id,
        "to":              req.fax_number,
    }


@app.post("/api/send-fax")
async def send_fax_api(
    req: SendFaxRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """
    POST /api/send-fax
    JWT認証 + 自院のファイルのみ FAX 送信可能。
    file_key は presign-upload で取得した R2 キー。
    fax_number はサーバ側で contacts から取得する（クライアント値は使わない）。
    送信結果は documents に source="fax_outbound" で記録。
    """
    jwt_token   = credentials.credentials
    user_id     = user.get("sub", "")
    hospital_id = _get_hospital_id(user_id, jwt_token)
    contact     = _get_fax_contact(req.contact_id, hospital_id, jwt_token)
    _assert_fax_file_key(req.file_key)
    return await _send_fax_impl(req, hospital_id, user_id, jwt_token, contact["fax_number"])


@app.post("/send-fax")
async def send_fax_compat(
    req: SendFaxRequest,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    user: dict = Depends(verify_jwt),
):
    """compat: Vite proxy 経由のローカル開発用（/api/send-fax と同じ処理）"""
    jwt_token   = credentials.credentials
    user_id     = user.get("sub", "")
    hospital_id = _get_hospital_id(user_id, jwt_token)
    contact     = _get_fax_contact(req.contact_id, hospital_id, jwt_token)
    _assert_fax_file_key(req.file_key)
    return await _send_fax_impl(req, hospital_id, user_id, jwt_token, contact["fax_number"])


@app.post("/webhook/cloudfax/outbound")
async def cloudfax_outbound_compat(request: Request):
    """compat: Vite proxy 経由のローカル開発用（/api/webhook/cloudfax/outbound と同じ処理）"""
    _verify_webhook_secret(request)

    try:
        payload_raw: dict = await request.json()
    except Exception:
        logger.error("[cloudfax/outbound] JSON パース失敗")
        raise HTTPException(status_code=400, detail="JSON パースに失敗しました")

    try:
        CloudFaxOutboundPayload.model_validate(payload_raw)
    except Exception as e:
        logger.error("[cloudfax/outbound] payload バリデーション失敗: %s", e)
        raise HTTPException(status_code=400, detail="payload のバリデーションに失敗しました")

    return await _cloudfax_outbound_impl(payload_raw)
